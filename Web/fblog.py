from selenium import webdriver
from selenium.webdriver.common.keys import Keys
import urllib.request
import numpy as np
import pandas as pd
from time import sleep
from docx import Document
from selenium.common.exceptions import StaleElementReferenceException


url = "https://www.facebook.com/"
chromedriver = r"D:\Programming Segate 25 Jul 2019\Prgramming_19-05-2019\WEB\chromedriver"
browser = webdriver.Chrome(chromedriver)

browser.get(url)
e_id = browser.find_element_by_id("email")
e_id.send_keys('phanindraparasshar@gmail.com')



link = 'https://www.facebook.com/sneha.srinivasan.716/friends'



def getFriend(Name,i):
    url_friend = "https://www.facebook.com/search/top/?q="+Name+"&epa=SEARCH_BOX"
    browser.get(url_friend)
    
    profile_xpath = '//*[@id="xt_uniq_'+str(i)+'"]/div/div/div[1]/div[1]/div[1]/a'
    friend = browser.find_element_by_xpath()
    
    friend.click()


    
def MyProfile():
    my_url = 'https://www.facebook.com/phanindra.parashar'
    browser.get(my_url)
    friends_st = browser.find_elements_by_partial_link_text('Friends')
    sleep(3)
    for i in friends_st:
        if 'Tag' not in  i.text:
            nFriends = int(i.text[7:])
            i.click()
            break
    browser.execute_script("window.scrollTo(0,400)")
    print("I have ",nFriends," friends")
    
    j = 0
    while j < 721:
        browser.execute_script("window.scrollTo(0,document.body.scrollHeight)")
        sleep(1)
        
        friends_elem = browser.find_elements_by_partial_link_text('friend')
        Furl = []
        for i in friends_elem:
            friend_url = i.get_attribute("href")
            if friend_url not in Furl:
                Furl.append(friend_url)
        j = len(Furl)
    return Furl

def get_li():
    L = browser.find_elements_by_tag_name("li")
    Val = []
    for i in L:
        id_val = i.get_attribute("id")
        if id_val != '':
            Val.append(id_val)
    return Val

def friend_profile(link):
    browser.get(link)
    browser.execute_script("window.scrollTo(0,-document.body.scrollHeight)")
    sleep(2)
    Name = browser.find_element_by_id("fb-timeline-cover-name").text
    Pics = browser.find_elements_by_tag_name('img')
    for i in Pics:
        if 'Profile Photo' in i.get_attribute('alt'):
            profile_pic = i.get_attribute('src')
            break
    
    
    print(Name)
    
    Look_up_list = ['Overview','Work','Places','Family','Details','Events']
    about_page = browser.find_element_by_partial_link_text("About").click()
    sleep(2)
    
    Text_about = ' '
    About_list = []
    for j in Look_up_list:
        Emt = j +  '   '
        
        Text_about = Text_about + j + '   : '
        about_page = browser.find_element_by_partial_link_text(j).click()
        sleep(2)
        about_ids = get_li()
        
        for i in about_ids:
            Text_about = Text_about + '. ' + browser.find_elements_by_id(i)[0].text
            Emt = Emt + '. ' + browser.find_elements_by_id(i)[0].text
        Text_about = Text_about + '          '
        About_list.append(Emt)
        
    return About_list,profile_pic,Name

def friend_time_line(link):
    browser.get(link)
    browser.execute_script("window.scrollTo(0,-document.body.scrollHeight)")
    Name = browser.find_element_by_id("fb-timeline-cover-name").text
    
    print(Name)
    
    time_line = browser.find_element_by_partial_link_text("Timeline").click()
    
    Text_post = ' '
    
    Posts = []
    j = 0
    while j < 24:
        j = j + 1
        
        posts = browser.find_elements_by_xpath("//*[contains(@id, 'jumper')]")
        for i in posts:
            if i not in Posts :
                Posts.append(i)
        sleep(1)
        browser.execute_script("window.scrollTo(0,document.body.scrollHeight)")
        
    only_post = []
    for i in Posts:
        data = i.text
        Text_post = Text_post + data + '                                    '
        only_post.append(data)
    
    
    return only_post

def download_image(image_src,name):
    urllib.request.urlretrieve(image_src,name+'.png')

def refine_post(post):
    schar = '·'
    b1 = post.replace('\n','. ')
    b2 = b1.replace(' Like. Comment. Share.','')
    b3 = b2.replace('Write a comment....','')
    b4 = b3.replace(schar,'')
    b5 = b4.replace('.  to.  ',' to ')
    b6 = b5.replace(' Comments.','')
    b7 = b6.replace('Write a comment...','')
    b8 = b7.replace('. Like.','')
    b9 = b8.replace('..','.')
    
    return b9
    
        

def make_docx(about,Post_List,name,profile_pic):
    schar = '·'
    download_image(profile_pic,name)
    document = Document()
    sleep(3)
    document.add_picture(name+'.png')
    document.add_heading(name, 0)
    
    Look_up_list = ['Overview','Work','Places','Family','Details','Events']
    
    for i in range(len(about)):
        heading_sub = Look_up_list[i]
        raw_text = about[i]
        i1 = raw_text.replace('\n','. ')
        i2 = i1.replace(heading_sub,'')
        for j in range(len(i2)):
            if i2[j].isalpha():
                break
        text_info = i2[j:]
        
        document.add_heading(heading_sub,level=1)
        document.add_paragraph(text_info)
    
    document.add_page_break()
    
    for i in range(len(Post_List)):
        content = refine_post(Post_List[i])
        document.add_paragraph(content, style='List Number')
        

    document.save(name+'.docx')

def collect_data(link):
    about,profile_pic,name = friend_profile(link)
    Post_List = friend_time_line(link)
    make_docx(about,Post_List,name,profile_pic)
    
    
    
    
    
    
        
        
        
        
        
    
    
    


    